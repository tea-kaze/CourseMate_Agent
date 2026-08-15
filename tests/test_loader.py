from __future__ import annotations

from pathlib import Path

import pytest

from coursemate.rag.loader import DocumentParseError, load_document, split_documents


def test_load_txt_document(tmp_path: Path):
    f = tmp_path / "note.txt"
    f.write_text("第一章：进程与线程。进程是资源分配的基本单位。", encoding="utf-8")
    docs = load_document(f)
    assert docs and "进程" in docs[0].page_content


def test_load_unsupported_type_raises(tmp_path: Path):
    f = tmp_path / "data.xlsx"
    f.write_bytes(b"not supported")
    with pytest.raises(DocumentParseError, match="不支持的文件类型"):
        load_document(f)


def test_load_docx_paragraph_and_table(tmp_path: Path):
    """.docx 应提取段落文本与表格内容。"""
    from docx import Document

    f = tmp_path / "course.docx"
    doc = Document()
    doc.add_paragraph("第一章：进程与线程。进程是资源分配的基本单位。")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "调度算法"
    table.cell(0, 1).text = "抢占式"
    table.cell(1, 0).text = "FCFS"
    table.cell(1, 1).text = "否"
    doc.save(f)

    docs = load_document(f)
    content = "\n".join(d.page_content for d in docs)
    assert "第一章：进程与线程" in content
    assert "调度算法 | 抢占式" in content
    assert "FCFS | 否" in content


def test_load_empty_docx_raises(tmp_path: Path):
    """空 .docx 应给出明确报错而不是崩溃。"""
    from docx import Document

    f = tmp_path / "empty.docx"
    doc = Document()
    doc.save(f)
    with pytest.raises(DocumentParseError, match="未提取到文本"):
        load_document(f)


def test_split_documents_keeps_content():
    from langchain_core.documents import Document as LCDocument

    text = "第一段。" * 300
    chunks = split_documents([LCDocument(page_content=text)])
    assert len(chunks) >= 2
    assert all(c.page_content for c in chunks)


def test_load_empty_pdf_raises(tmp_path: Path):
    """空/无文本 PDF 应给出明确报错而不是崩溃。"""
    from pypdf import PdfWriter

    f = tmp_path / "empty.pdf"
    writer = PdfWriter()
    with f.open("wb") as fh:
        writer.write(fh)
    with pytest.raises(DocumentParseError, match="未提取到文本"):
        load_document(f)
